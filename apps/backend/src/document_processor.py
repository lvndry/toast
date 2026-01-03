"""Document processing service for uploaded documents with OCR, classification, and summarization."""

import importlib
import json
from io import BytesIO
from typing import Any

import pdfplumber
from docx import Document as DocxDocument
from dotenv import load_dotenv
from pydantic import BaseModel

from src.core.logging import get_logger
from src.llm import SupportedModel, acompletion_with_fallback
from src.models.document import Document, DocumentAnalysis
from src.summarizer import summarize_document

load_dotenv()
logger = get_logger(__name__)


class DocumentProcessingResult(BaseModel):
    """Result of document processing."""

    success: bool
    is_legal_document: bool
    document: Document | None = None
    analysis: DocumentAnalysis | None = None
    error_message: str | None = None


class DocumentProcessor:
    """Process uploaded documents with OCR, classification, and summarization."""

    def __init__(
        self,
        model: SupportedModel | None = None,
        max_content_length: int = 5000,
        enable_binary_parsing: bool = False,
        prefer_tika: bool = False,
        prefer_pdfminer: bool = False,
    ):
        """Initialize the document processor.

        Args:
            enable_binary_parsing: If True, attempt to extract text from binary files (PDF/Office) using
                optional parsers (pdfminer / tika) when available.
            prefer_tika: If True and tika is available, prefer Tika for broad binary parsing (Office, PDF).
            prefer_pdfminer: If True and pdfminer is available, use it as a fallback for PDFs when pdfplumber fails.
        """
        # Store model_name for potential future use, but we'll use fallback system
        self.model_name = model
        self.max_content_length = max_content_length

        # Binary parsing configuration (opt-in)
        self.enable_binary_parsing = enable_binary_parsing
        self.prefer_tika = prefer_tika
        self.prefer_pdfminer = prefer_pdfminer

        # Document type categories for classification
        self.categories = [
            "privacy_policy",
            "terms_of_service",
            "cookie_policy",
            "terms_and_conditions",
            "data_processing_agreement",
            "gdpr_policy",
            "copyright_policy",
            "safety_policy",
            "other",
        ]

    async def process_document(
        self, file_content: bytes, filename: str, content_type: str, company_id: str
    ) -> DocumentProcessingResult:
        """
        Process an uploaded document through the complete pipeline.

        Args:
            file_content: Raw file content
            filename: Original filename
            content_type: MIME type of the file
            company_id: ID of the company/conversation

        Returns:
            DocumentProcessingResult with processing results
        """
        try:
            # Step 1: Extract text from document (OCR for PDFs, direct text for others)
            text_content = await self._extract_text(file_content, filename, content_type)
            if not text_content:
                return DocumentProcessingResult(
                    success=False,
                    is_legal_document=False,
                    error_message="Failed to extract text from document",
                )

            # Step 2: Classify the document
            classification = await self._classify_document(text_content, filename)

            # Step 3: Check if it's a legal document
            if not classification.get("is_legal_document", False):
                return DocumentProcessingResult(
                    success=False,
                    is_legal_document=False,
                    error_message="Document is not classified as a legal document",
                )

            # Step 4: Create document object
            document = Document(
                url=f"uploaded://{filename}",
                title=filename,
                company_id=company_id,
                doc_type=classification.get("classification", "other"),
                markdown=text_content,
                text=text_content,
                metadata={
                    "filename": filename,
                    "content_type": content_type,
                    "size": len(file_content),
                    "classification": classification,
                },
            )

            # Step 5: Generate summary using existing summarizer
            analysis = await summarize_document(document)
            if analysis:
                document.analysis = analysis

            return DocumentProcessingResult(
                success=True,
                is_legal_document=True,
                document=document,
                analysis=analysis,
            )

        except Exception as e:
            logger.error(f"Error processing document {filename}: {str(e)}")
            return DocumentProcessingResult(
                success=False,
                is_legal_document=False,
                error_message=f"Processing error: {str(e)}",
            )

    async def _extract_text(
        self, file_content: bytes, filename: str, content_type: str
    ) -> str | None:
        """Extract text from uploaded document."""
        try:
            # Handle different file types
            if content_type == "application/pdf":
                return await self._extract_text_from_pdf(file_content)
            elif content_type in ["text/plain", "text/markdown"]:
                return file_content.decode("utf-8", errors="ignore")
            elif content_type in [
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ]:
                return await self._extract_text_from_docx(file_content)
            else:
                # Try to decode as text for unknown types
                return file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            return None

    async def _extract_text_from_pdf(self, file_content: bytes) -> str | None:
        """Extract text from PDF using pdfplumber, with optional fallbacks (pdfminer / tika).

        This method prefers pdfplumber for text-based PDFs (fast and reliable). If that
        fails to produce text and binary parsing is enabled, it will optionally try
        pdfminer (if available) and then Tika (if configured and available).
        """
        try:
            # Primary strategy: pdfplumber
            with BytesIO(file_content) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    text_parts = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                    if text_parts:
                        return "\n".join(text_parts)
                    else:
                        # If no text extracted, the PDF might be image-based
                        logger.warning("No text extracted from PDF - may need OCR or other parser")

            # If binary parsing is not enabled, return
            if not self.enable_binary_parsing:
                return None

            # Optional fallback: pdfminer.six (if available and requested)
            if self.prefer_pdfminer:
                try:
                    from pdfminer.high_level import extract_text as pdfminer_extract_text

                    with BytesIO(file_content) as fp:
                        text = pdfminer_extract_text(fp)
                        if text and text.strip():
                            return text
                except Exception as e:
                    logger.warning(f"pdfminer fallback failed: {e}")

            # Optional: Tika (broad support for PDF and Office). Tika may not be installed.
            if self.prefer_tika:
                try:
                    tika_module = importlib.import_module("tika")
                    tika_parser = tika_module.parser

                    parsed = tika_parser.from_buffer(file_content)
                    text = parsed.get("content") if isinstance(parsed, dict) else None
                    if text and text.strip():
                        return text
                except Exception as e:
                    logger.warning(f"Tika parsing failed or not available: {e}")

            return None

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return None

    async def _extract_text_from_docx(self, file_content: bytes) -> str | None:
        """Extract text from DOCX file using python-docx, or fallback to Tika when enabled.

        Tika provides broader support for legacy Office formats and binary extraction but it is
        optional and not installed by default. This method will only try Tika when
        `self.enable_binary_parsing` and `self.prefer_tika` are True.
        """
        try:
            # Primary: python-docx for modern .docx
            with BytesIO(file_content) as docx_file:
                doc = DocxDocument(docx_file)
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                if text_parts:
                    return "\n".join(text_parts)
                else:
                    # Fallthrough to optional Tika if configured
                    logger.debug(
                        "No text extracted from DOCX by python-docx; attempting Tika if enabled"
                    )

            if not (self.enable_binary_parsing and self.prefer_tika):
                return None

            try:
                tika_module = importlib.import_module("tika")
                tika_parser = tika_module.parser

                parsed = tika_parser.from_buffer(file_content)
                text = parsed.get("content") if isinstance(parsed, dict) else None
                if text and text.strip():
                    return text
            except Exception as e:
                logger.warning(f"Tika parsing for DOCX failed or not available: {e}")

            return None
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return None

    async def _classify_document(self, text_content: str, filename: str) -> dict[str, Any]:
        """Classify the document using LLM."""
        # Truncate content if too long
        if len(text_content) > self.max_content_length:
            text_content = text_content[: self.max_content_length]

        prompt = f"""Analyze this document and classify it as a legal document.

Document filename: {filename}
Document content:
{text_content}

Please return a JSON object with the following fields:

- classification: the most appropriate category from this list: {self.categories}
- classification_justification: a brief explanation of why this category was selected
- is_legal_document: a boolean. This should be True only if the document contains substantive legal text (e.g., terms of service, privacy policy, data protection policy, etc.)
- is_legal_document_justification: a short rationale for your legal classification decision

Example output:
{{
  "classification": "privacy_policy",
  "classification_justification": "The document outlines data collection practices, user rights, and privacy protections consistent with a privacy policy.",
  "is_legal_document": true,
  "is_legal_document_justification": "The document contains substantive legal text defining privacy rights, data handling obligations, and user consent mechanisms."
}}

Use caution: If the content appears incomplete, vague, or primarily promotional, treat it with skepticism and prefer "other" unless clear evidence suggests a more specific classification."""

        system_prompt = """You are a legal document classifier. Identify substantive legal content and categorize accurately."""

        try:
            response = await acompletion_with_fallback(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                response_format={"type": "json_object"},
            )

            choice = response.choices[0]
            if not hasattr(choice, "message"):
                raise ValueError("Unexpected response format: missing message attribute")
            message = choice.message  # type: ignore[attr-defined]
            if not message:
                raise ValueError("Unexpected response format: message is None")
            content = message.content  # type: ignore[attr-defined]
            if not content:
                raise ValueError("Empty response from LLM")

            result: dict[str, Any] = json.loads(content)
            logger.debug(f"Document classification result: {result}")
            return result

        except Exception as e:
            logger.warning(f"Document classification failed: {e}")
            return {
                "classification": "other",
                "classification_justification": f"Classification failed: {e}",
                "is_legal_document": False,
                "is_legal_document_justification": "Could not analyze due to error",
            }
